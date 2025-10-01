"""RAG Engine using LlamaIndex and Qdrant for document retrieval and analysis."""

from datetime import datetime
import logging
from logging import Logger
from pathlib import Path
from typing import Any, List, Optional, Tuple

from llama_index.core import Document, Settings, SimpleDirectoryReader, StorageContext, VectorStoreIndex
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms.openai import OpenAI
from llama_index.vector_stores.qdrant import QdrantVectorStore
import numpy as np
from qdrant_client.models import DatetimeRange, Distance, FieldCondition, Filter, VectorParams

from config.settings import settings
from services.audio_overview_service import clean_markdown
from services.format_helper import format_analysis_result
from services.prompt_router import choose_prompt
from services.query_cache import QueryCache
from src.domain.entities.tenant_context import TenantContext
from src.infrastructure.performance.connection_pool import get_qdrant_pool, get_query_optimizer

try:
    from src.application.services.enterprise_orchestrator import EnterpriseOrchestrator, EnterpriseQuery

    ENTERPRISE_AVAILABLE = True
except ImportError:
    Logger.warning("Enterprise orchestrator not available")
    ENTERPRISE_AVAILABLE = False
    EnterpriseOrchestrator = None
    EnterpriseQuery = None

try:
    from src.application.services.hyde_retriever import HyDEConfig, HyDEQueryEngine
    from src.application.services.streaming_rag import StreamingChunk, StreamingRAGEngine

    ADVANCED_FEATURES_AVAILABLE = True
except ImportError:
    Logger.warning("Advanced features (Streaming/HyDE) not available")
    ADVANCED_FEATURES_AVAILABLE = False
    StreamingRAGEngine = None
    HyDEQueryEngine = None

try:
    from services.contextual_retrieval_service import get_contextual_retrieval_service
    from services.reranking_service import get_reranking_service

    QUALITY_FEATURES_AVAILABLE = True
except ImportError:
    logger.warning("Quality enhancement features (Reranking/Contextual) not available")
    QUALITY_FEATURES_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class RAGEngine:
    """RAG engine for document indexing and retrieval using LlamaIndex and Qdrant."""

    def __init__(self, tenant_context: Optional[TenantContext] = None):
        """Initialize RAG engine with Qdrant and OpenAI, optionally with tenant context."""
        self.client = None
        self.vector_store = None
        self.index = None
        self.tenant_context = tenant_context
        self.collection_name = self._get_tenant_collection_name()
        # Initialize query cache if enabled (with tenant namespace if multi-tenant)
        cache_namespace = f"tenant_{tenant_context.tenant_id}" if tenant_context else None
        self.query_cache = (
            QueryCache(ttl_seconds=3600, namespace=cache_namespace) if settings.rag_enable_caching else None
        )
        # Initialize enterprise orchestrator
        self.enterprise_orchestrator = None
        # Initialize quality enhancement services
        self.reranking_service = None
        self.contextual_service = None
        # Use connection pool for Qdrant
        self.connection_pool = get_qdrant_pool()
        self.query_optimizer = get_query_optimizer()
        self._initialize_components()

    def _initialize_components(self):
        """Initialize Qdrant client, vector store, and global settings."""
        try:
            # Configure global LlamaIndex settings
            Settings.llm = OpenAI(
                model=settings.llm_model,
                temperature=settings.temperature,
                max_tokens=settings.max_tokens,
                api_key=settings.openai_api_key,
            )
            Settings.embed_model = OpenAIEmbedding(model=settings.embedding_model, api_key=settings.openai_api_key)
            Settings.chunk_size = settings.chunk_size
            Settings.chunk_overlap = settings.chunk_overlap

            # Use connection pool for Qdrant client
            with self.connection_pool.get_connection() as client:
                self.client = client

            # Create or recreate collection
            self._setup_collection()

            # Initialize vector store
            self.vector_store = QdrantVectorStore(client=self.client, collection_name=self.collection_name)

            # Initialize or load existing index
            self._initialize_index()

            # Initialize advanced features (Streaming and HyDE)
            self.streaming_engine = None
            self.hyde_engine = None

            if ADVANCED_FEATURES_AVAILABLE and self.index:
                try:
                    # Initialize streaming engine
                    query_engine = self.index.as_query_engine()
                    self.streaming_engine = StreamingRAGEngine(
                        query_engine=query_engine, llm=Settings.llm, chunk_size=10, stream_delay=0.01
                    )

                    # Initialize HyDE engine
                    hyde_config = HyDEConfig(num_hypothetical_docs=3, temperature=0.0, include_original_query=True)
                    self.hyde_engine = HyDEQueryEngine(
                        index=self.index,
                        llm=Settings.llm,
                        hyde_config=hyde_config,
                        domain="financial",  # Set to financial for business documents
                    )
                    logger.info("Advanced features (Streaming/HyDE) initialized")
                except Exception as e:
                    logger.warning(f"Could not initialize advanced features: {str(e)}")

            # Initialize enterprise orchestrator after RAG components are ready
            if ENTERPRISE_AVAILABLE:
                self.enterprise_orchestrator = EnterpriseOrchestrator(
                    rag_engine=self,
                    csv_analyzer=None,  # Will be set externally if needed
                    openai_api_key=settings.openai_api_key,  # Pass API key for embeddings
                )
            else:
                self.enterprise_orchestrator = None

            # Initialize quality enhancement services
            if QUALITY_FEATURES_AVAILABLE:
                try:
                    # Initialize reranking service with default model
                    self.reranking_service = get_reranking_service(model_name="default")
                    logger.info("Reranking service initialized")

                    # Initialize contextual retrieval service
                    self.contextual_service = get_contextual_retrieval_service(window_size=1)
                    logger.info("Contextual retrieval service initialized")
                except Exception as e:
                    logger.warning(f"Could not initialize quality enhancement services: {str(e)}")
                    self.reranking_service = None
                    self.contextual_service = None

            tenant_info = f" for tenant {self.tenant_context.tenant_id}" if self.tenant_context else ""
            quality_info = " with quality enhancements" if self.reranking_service and self.contextual_service else ""
            logger.info(f"RAG Engine initialized successfully with enterprise orchestrator{quality_info}{tenant_info}")

        except Exception as e:
            logger.error(f"Error initializing RAG Engine: {str(e)}")
            raise

    def _get_tenant_collection_name(self) -> str:
        """Get collection name based on tenant context."""
        if self.tenant_context:
            return f"tenant_{self.tenant_context.tenant_id}_docs"
        return settings.qdrant_collection_name

    def _setup_collection(self):
        """Setup Qdrant collection with proper configuration."""
        try:
            # Check if collection exists
            collections = self.client.get_collections().collections
            collection_exists = any(c.name == self.collection_name for c in collections)

            if collection_exists:
                logger.info(f"Collection {self.collection_name} already exists")
            else:
                # Create new collection with proper vector size for OpenAI embeddings
                vector_size = 1536  # OpenAI text-embedding-3-small dimension
                try:
                    self.client.create_collection(
                        collection_name=self.collection_name,
                        vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
                    )
                    logger.info(f"Created new collection: {self.collection_name}")
                except Exception as create_error:
                    # If collection creation fails due to orphaned data, try to delete and recreate
                    if (
                        "already exists" in str(create_error).lower()
                        or "data already exists" in str(create_error).lower()
                    ):
                        logger.warning(
                            f"Collection data exists but not accessible. Attempting to clean and recreate: {create_error}"
                        )
                        try:
                            # Try to delete the collection first
                            self.client.delete_collection(collection_name=self.collection_name)
                            logger.info(f"Deleted orphaned collection data for {self.collection_name}")
                        except Exception as delete_error:
                            logger.warning(f"Could not delete orphaned collection: {delete_error}")

                        # Now try to create again
                        self.client.create_collection(
                            collection_name=self.collection_name,
                            vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
                        )
                        logger.info(f"Successfully recreated collection: {self.collection_name}")
                    else:
                        raise create_error

        except Exception as e:
            logger.error(f"Error setting up collection: {str(e)}")
            raise

    def _initialize_index(self):
        """Initialize or load existing index."""
        try:
            storage_context = StorageContext.from_defaults(vector_store=self.vector_store)

            # Check if index exists in vector store
            collection_info = self.client.get_collection(self.collection_name)
            if collection_info.points_count > 0:
                # Load existing index
                self.index = VectorStoreIndex.from_vector_store(self.vector_store)
                logger.info(f"Loaded existing index with {collection_info.points_count} vectors")
            else:
                # Create new empty index
                self.index = VectorStoreIndex([], storage_context=storage_context)
                logger.info("Created new empty index")

        except Exception as e:
            logger.error(f"Error initializing index: {str(e)}")
            self.index = VectorStoreIndex(
                [], storage_context=StorageContext.from_defaults(vector_store=self.vector_store)
            )

    def parse_insert_docs(
        self,
        file_paths: list[str],
        metadata: Optional[dict[str, Any]] = None,
        original_names: Optional[list[str]] = None,
        permanent_paths: Optional[list[str]] = None,
    ) -> dict[str, Any]:
        from services import csv_analyzer

        results = {
            "uploaded_files": [],
            "failed_files": [],
            "total_chunks": 0,
            "errors": [],
        }

        for i, file_path in enumerate(file_paths):
            try:
                path = Path(file_path)
                if not path.exists():
                    results["failed_files"].append(file_path)
                    results["errors"].append(f"File not found: {file_path}")
                    continue
                # use original filename if provided
                display_name = original_names[i] if original_names and i < len(original_names) else path.name

                # load document based on file type
                if path.suffix.lower() == ".pdf":
                    documents = self._load_pdf(file_path, metadata)
                elif path.suffix.lower() in [".txt", ".md"]:
                    documents = self._load_text(file_path, metadata)
                elif path.suffix.lower() in [".docx", ".doc"]:
                    documents = self._load_docx(file_path, metadata)
                elif path.suffix.lower() == ".json":
                    # JSON files are loaded with enhanced metadata detection
                    documents = self._load_json(file_path, metadata)
                elif path.suffix.lower() == ".csv" and csv_analyzer:
                    # CSV files are analyzed and converted to structured documents
                    documents = self._load_csv_with_analysis(file_path, csv_analyzer, metadata)
                elif path.suffix.lower() == ".csv":
                    # Basic CSV loading without analysis
                    documents = self._load_csv_basic(file_path, metadata)
                else:
                    documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
                for doc in documents:
                    doc.metadata = doc.metadata or {}
                    doc.metadata.update(
                        {
                            "source": display_name,
                            "indexed_at": datetime.now().isoformat(),
                            "file_types": path.suffix.lower(),
                        }
                    )
                    if metadata:
                        doc.metadata.update(metadata)
                    parser = SimpleNodeParser.from_defaults()
                    nodes = parser.get_nodes_from_documents(documents)
                    self.index.insert_nodes(nodes)

                    results["uploaded_files"].append(file_path)
                    results["total_chunks"] += len(nodes)
                    logger.info(f"Uploaded {file_path} : {len(nodes)} chunks")
            except Exception as e:
                results["failed_files"].append(file_path)
                results["errors"].append(f"Error uploading {file_path}:{str(e)}")
                logger.info(f"Uploaded {file_path} : {str(e)}")
        return results

    def index_documents(
        self,
        file_paths: list[str],
        metadata: Optional[dict[str, Any]] = None,
        original_names: Optional[list[str]] = None,
        permanent_paths: Optional[list[str]] = None,
        force_prompt_type: Optional[str] = None,
        csv_analyzer=None,
    ) -> dict[str, Any]:
        """Index documents from file paths."""
        results = {
            "indexed_files": [],
            "failed_files": [],
            "total_chunks": 0,
            "errors": [],
            "document_analyses": {},  # Store automatic analyses
        }

        for i, file_path in enumerate(file_paths):
            try:
                path = Path(file_path)
                if not path.exists():
                    results["failed_files"].append(file_path)
                    results["errors"].append(f"File not found: {file_path}")
                    continue

                # Use original filename if provided, otherwise use current filename
                display_name = original_names[i] if original_names and i < len(original_names) else path.name
                # Get permanent path for PDF viewing
                permanent_path = permanent_paths[i] if permanent_paths and i < len(permanent_paths) else None

                # Load document based on file type
                if path.suffix.lower() == ".pdf":
                    documents = self._load_pdf(file_path, metadata)
                elif path.suffix.lower() in [".txt", ".md"]:
                    documents = self._load_text(file_path, metadata)
                elif path.suffix.lower() in [".docx", ".doc"]:
                    documents = self._load_docx(file_path, metadata)
                elif path.suffix.lower() == ".json":
                    # JSON files are loaded with enhanced metadata detection
                    documents = self._load_json(file_path, metadata)
                elif path.suffix.lower() == ".csv" and csv_analyzer:
                    # CSV files are analyzed and converted to structured documents
                    documents = self._load_csv_with_analysis(file_path, csv_analyzer, metadata)
                elif path.suffix.lower() == ".csv":
                    # Basic CSV loading without analysis
                    documents = self._load_csv_basic(file_path, metadata)
                elif path.suffix.lower() in [".xls", ".xlsx"]:
                    # Excel files are converted to documents
                    documents = self._load_excel(file_path, metadata)
                elif path.suffix.lower() in [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"]:
                    # Image files are processed with OCR
                    documents = self._load_image(file_path, metadata)
                else:
                    documents = SimpleDirectoryReader(input_files=[file_path]).load_data()

                # Add metadata to documents
                for doc in documents:
                    doc.metadata = doc.metadata or {}
                    doc_metadata = {
                        "source": display_name,  # Nome file leggibile
                        "indexed_at": datetime.now().isoformat(),
                        "file_type": path.suffix.lower(),
                        "document_size": len(doc.text) if hasattr(doc, "text") else 0,
                    }
                    # Add permanent path for PDF viewing (only for PDFs)
                    if permanent_path and path.suffix.lower() == ".pdf":
                        doc_metadata["pdf_path"] = permanent_path

                    doc.metadata.update(doc_metadata)
                    if metadata:
                        doc.metadata.update(metadata)

                # Parse documents into nodes
                parser = SimpleNodeParser.from_defaults()
                nodes = parser.get_nodes_from_documents(documents)

                # Add nodes to index
                self.index.insert_nodes(nodes)

                # Generate automatic analysis of the document content
                if documents:
                    full_text = "\n".join([doc.text for doc in documents])

                    # Cache the document text for potential re-analysis
                    if not hasattr(self, "_last_document_texts"):
                        self._last_document_texts = {}
                    self._last_document_texts[display_name] = full_text

                    analysis = self.analyze_document_content(full_text, display_name, force_prompt_type)
                    results["document_analyses"][display_name] = analysis

                results["indexed_files"].append(file_path)
                results["total_chunks"] += len(nodes)
                logger.info(f"Indexed {file_path}: {len(nodes)} chunks")

            except Exception as e:
                results["failed_files"].append(file_path)
                results["errors"].append(f"Error indexing {file_path}: {str(e)}")
                logger.error(f"Error indexing {file_path}: {str(e)}")

        return results

    def clean_metadata_paths(self) -> bool:
        """Remove temporary paths from existing document metadata."""
        try:
            logger.info("Cleaning metadata paths from existing documents...")
            # Use the same logic as delete_documents to ensure clean slate
            return self.delete_documents("*")
        except Exception as e:
            logger.error(f"Error cleaning metadata: {str(e)}")
            return False

    def analyze_document_content(
        self, document_text: str, file_name: str, force_prompt_type: Optional[str] = None
    ) -> str:
        """Generate automatic analysis of document content using specialized prompts."""
        try:
            from openai import OpenAI

            client = OpenAI(api_key=settings.openai_api_key)

            # Truncate text if too long (keep first 8000 chars for analysis)
            analysis_text = document_text[:8000] if len(document_text) > 8000 else document_text

            logger.debug(f"Document text length: {len(document_text)}, Analysis text length: {len(analysis_text)}")
            logger.debug(f"Analysis text preview: {analysis_text[:200]}...")

            # Use prompt router to select the best prompt (or force a specific one)
            if force_prompt_type and force_prompt_type != "automatico":
                # Force a specific prompt type
                from services.prompt_router import PROMPT_GENERAL, ROUTER

                if force_prompt_type == "generale":
                    # Handle general prompt explicitly
                    prompt_name = "generale"
                    prompt_text = PROMPT_GENERAL(file_name, analysis_text)
                    debug_info = {"forced": True, "type": "generale"}
                    logger.info(f"Forcing general prompt for document '{file_name}'")
                elif force_prompt_type in ROUTER:
                    # Use specialized prompt from router
                    prompt_name = force_prompt_type
                    prompt_text = ROUTER[force_prompt_type].builder(file_name, analysis_text)
                    debug_info = {"forced": True, "type": force_prompt_type}
                    logger.info(f"Forcing prompt type '{prompt_name}' for document '{file_name}'")
                else:
                    # Unknown prompt type, fallback to general
                    prompt_name = "generale"
                    prompt_text = PROMPT_GENERAL(file_name, analysis_text)
                    debug_info = {"forced": True, "type": "generale", "fallback": True}
                    logger.warning(
                        f"Unknown prompt type '{force_prompt_type}', using general prompt for document '{file_name}'"
                    )
            else:
                # Use automatic selection
                prompt_name, prompt_text, debug_info = choose_prompt(file_name, analysis_text)
                logger.info(f"Auto-selected prompt type '{prompt_name}' for document '{file_name}'")
            logger.debug(f"Prompt selection debug info: {debug_info}")

            response = client.chat.completions.create(
                model=settings.llm_model,
                messages=[
                    {
                        "role": "system",
                        "content": "Sei un analista aziendale esperto che crea analisi professionali di documenti. Rispondi sempre in italiano perfetto e in modo strutturato e chiaro. Segui esattamente il formato richiesto nel prompt.",
                    },
                    {"role": "user", "content": prompt_text},
                ],
                temperature=0.0,  # Deterministic output
                max_tokens=4000,  # Increased for complex JSON + summary format (especially scadenzario)
            )

            analysis_result = response.choices[0].message.content

            # Add prompt type metadata to result for raw format
            raw_result = f"[Analisi tipo: {prompt_name.upper()}]\n\n{analysis_result}"

            # Format the result for better readability
            formatted_result = format_analysis_result(raw_result)

            return formatted_result

        except Exception as e:
            logger.error(f"Error analyzing document content: {str(e)}")
            return "Analisi automatica non disponibile per questo documento. Contenuto indicizzato correttamente per le ricerche."

    def reanalyze_documents_with_prompt(self, force_prompt_type: str) -> dict[str, str]:
        """Re-analyze existing documents with a specific prompt type."""
        try:
            # Check if we have existing analyses in memory that can be re-processed
            if hasattr(self, "_last_document_texts") and self._last_document_texts:
                logger.info("Using cached document texts for re-analysis")
                results = {}
                for source, text in self._last_document_texts.items():
                    if text and text.strip():
                        logger.info(f"Re-analyzing cached document '{source}': {len(text)} characters")
                        analysis = self.analyze_document_content(text, source, force_prompt_type)
                        results[source] = analysis
                    else:
                        results[source] = (
                            f"## Errore di Ri-analisi\n\nNessun contenuto disponibile per il documento '{source}'."
                        )
                return results

            # Fallback: try to get content from vector store (less reliable)
            if not self.index:
                return {"error": "Nessun documento indicizzato trovato"}

            results = {}

            # Get all documents from the vector store
            collection_info = self.client.get_collection(self.collection_name)
            if collection_info.points_count == 0:
                return {"error": "Nessun documento nel database vettoriale"}

            # Try to retrieve documents using a broad query
            logger.info("Attempting to retrieve document content via similarity search")
            query_engine = self.index.as_query_engine(similarity_top_k=50, response_mode="no_text")

            # Use a generic query to get document nodes
            try:
                response = query_engine.query("contenuto documento analisi")

                if hasattr(response, "source_nodes") and response.source_nodes:
                    documents_content = {}

                    for node in response.source_nodes:
                        source = node.node.metadata.get("source", "Unknown")
                        text = node.node.text or node.node.get_content()

                        if source not in documents_content:
                            documents_content[source] = []
                        documents_content[source].append(text)

                    logger.info(f"Retrieved content from {len(documents_content)} documents via query")

                    # Re-analyze each document
                    for source, text_chunks in documents_content.items():
                        try:
                            # Combine all chunks for this document
                            full_text = "\n\n".join(text_chunks)

                            logger.info(
                                f"Re-analyzing document '{source}': {len(text_chunks)} chunks, {len(full_text)} total characters"
                            )

                            # Check if we have actual content
                            if not full_text.strip() or len(full_text.strip()) < 50:
                                logger.warning(f"Insufficient content for document '{source}' ({len(full_text)} chars)")
                                results[source] = f"""## Documento Non Disponibile per Ri-analisi

**Documento:** {source}

**Problema:** Il contenuto del documento non è disponibile per la ri-analisi. Questo può accadere quando:
- Il documento originale conteneva principalmente immagini
- Si è verificato un errore durante l'estrazione del testo
- Il documento è stato indicizzato ma il contenuto non è completamente recuperabile

**Soluzione:** Per ri-analizzare questo documento:
1. Ricarica il documento originale
2. Assicurati che contenga testo selezionabile
3. Considera l'uso di OCR se il documento è basato su immagini"""
                                continue

                            # Re-analyze with the forced prompt
                            analysis = self.analyze_document_content(full_text, source, force_prompt_type)
                            results[source] = analysis

                        except Exception as e:
                            logger.error(f"Error re-analyzing document {source}: {str(e)}")
                            results[source] = f"Errore nella ri-analisi: {str(e)}"

                    return results
                else:
                    return {"error": "Impossibile recuperare il contenuto dei documenti dal database vettoriale"}

            except Exception as query_error:
                logger.error(f"Error querying for document content: {str(query_error)}")
                return {"error": f"Errore nel recupero del contenuto: {str(query_error)}"}

        except Exception as e:
            logger.error(f"Error in reanalyze_documents_with_prompt: {str(e)}")
            return {"error": f"Errore nella ri-analisi: {str(e)}"}

    def _load_pdf(self, file_path: str, metadata: Optional[dict[str, Any]] = None) -> list[Document]:
        """Load and parse PDF documents with enhanced text extraction."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            documents = []
            total_extracted_text = ""

            logger.info(f"Processing PDF: {file_path} with {len(reader.pages)} pages")

            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    page_text_length = len(text.strip())
                    total_extracted_text += text

                    logger.debug(f"Page {i + 1}: extracted {page_text_length} characters")

                    if text.strip():
                        doc = Document(
                            text=text,
                            metadata={
                                "page": i + 1,
                                "total_pages": len(reader.pages),
                                "source": file_path,
                                "page_text_length": page_text_length,
                            },
                        )
                        if metadata:
                            doc.metadata.update(metadata)
                        documents.append(doc)
                    else:
                        logger.warning(f"Page {i + 1} in {file_path} has no extractable text")

                except Exception as page_error:
                    logger.error(f"Error extracting text from page {i + 1} in {file_path}: {str(page_error)}")
                    continue

            total_text_length = len(total_extracted_text.strip())
            logger.info(
                f"PDF processing complete: {len(documents)} pages with text, {total_text_length} total characters"
            )

            # If we extracted very little text, try alternative approach
            if total_text_length < 100:
                logger.warning(
                    f"Very little text extracted ({total_text_length} chars). PDF might be image-based or have extraction issues."
                )

                # Create a placeholder document with a warning
                if not documents:
                    warning_doc = Document(
                        text=f"[ATTENZIONE] Il documento PDF '{file_path}' potrebbe contenere principalmente immagini o avere problemi di estrazione testo. "
                        f"Sono stati estratti solo {total_text_length} caratteri da {len(reader.pages)} pagine. "
                        f"Per un'analisi completa, si consiglia di utilizzare un documento con testo selezionabile o di convertire "
                        f"le immagini in testo utilizzando OCR.",
                        metadata={
                            "page": 1,
                            "total_pages": len(reader.pages),
                            "source": file_path,
                            "extraction_warning": True,
                            "total_text_length": total_text_length,
                        },
                    )
                    if metadata:
                        warning_doc.metadata.update(metadata)
                    documents.append(warning_doc)

            return documents

        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise

    def _load_text(self, file_path: str, metadata: Optional[dict[str, Any]] = None) -> list[Document]:
        """Load text or markdown files."""
        try:
            with open(file_path, encoding="utf-8") as f:
                content = f.read()

            doc = Document(text=content, metadata={"source": file_path})
            if metadata:
                doc.metadata.update(metadata)

            return [doc]

        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {str(e)}")
            raise

    def _load_json(self, file_path: str, metadata: Optional[dict[str, Any]] = None) -> list[Document]:
        """Load JSON files with enhanced metadata detection."""
        try:
            import json

            with open(file_path, encoding="utf-8") as f:
                content = f.read()

            # Enhanced metadata for JSON files
            enhanced_metadata = {"source": file_path, "file_type": "json"}

            # Try to detect document type from JSON content
            try:
                json_data = json.loads(content)
                if isinstance(json_data, dict):
                    # Detect scadenzario/AR documents
                    scadenzario_keys = [
                        "aging_bucket",
                        "past_due",
                        "dso",
                        "dpd",
                        "crediti_lordi",
                        "crediti_netto",
                        "fondo_svalutazione",
                        "totale_scaduto",
                        "qualita_crediti",
                        "concentrazione_rischio",
                        "turnover_crediti",
                    ]

                    # Count matching keys
                    key_matches = sum(1 for key in scadenzario_keys if key in str(json_data))

                    if key_matches >= 3:  # If we have 3+ scadenzario keys
                        enhanced_metadata.update(
                            {
                                "document_type": "scadenzario",
                                "content_category": "accounts_receivable",
                                "analysis_type": "credit_management",
                                "scadenzario_keys_detected": key_matches,
                            }
                        )
                        logger.info(f"Detected scadenzario JSON with {key_matches} matching keys: {file_path}")

                    # Detect other financial document types
                    elif any(key in str(json_data) for key in ["ricavi", "ebitda", "bilancio", "conto_economico"]):
                        enhanced_metadata.update(
                            {
                                "document_type": "bilancio",
                                "content_category": "financial_statements",
                                "analysis_type": "financial_analysis",
                            }
                        )
                    elif any(key in str(json_data) for key in ["fatturato", "vendite", "sales"]):
                        enhanced_metadata.update(
                            {
                                "document_type": "fatturato",
                                "content_category": "sales_revenue",
                                "analysis_type": "sales_analysis",
                            }
                        )
                    else:
                        enhanced_metadata.update(
                            {
                                "document_type": "general",
                                "content_category": "structured_data",
                                "analysis_type": "general_analysis",
                            }
                        )

            except json.JSONDecodeError:
                # If JSON is malformed, treat as general text
                enhanced_metadata.update(
                    {
                        "document_type": "general",
                        "content_category": "text",
                        "analysis_type": "general_analysis",
                        "json_parse_error": True,
                    }
                )

            # Apply any additional metadata
            if metadata:
                enhanced_metadata.update(metadata)

            doc = Document(text=content, metadata=enhanced_metadata)
            return [doc]

        except Exception as e:
            logger.error(f"Error loading JSON file {file_path}: {str(e)}")
            raise

    def _load_docx(self, file_path: str, metadata: Optional[dict[str, Any]] = None) -> list[Document]:
        """Load Word documents."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            document = Document(text="\n".join(full_text), metadata={"source": file_path})
            if metadata:
                document.metadata.update(metadata)

            return [document]

        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise

    def query(
        self,
        query_text: str,
        top_k: int = 3,
        filters: Optional[dict[str, Any]] = None,
        analysis_type: Optional[str] = None,
    ) -> dict[str, Any]:
        """Query the indexed documents with optional specialized analysis."""
        try:
            if not self.index:
                return {"answer": "Nessun documento è stato ancora indicizzato.", "sources": [], "confidence": 0}

            # Check cache first if enabled
            if self.query_cache:
                cached_result = self.query_cache.get(query_text, top_k, analysis_type)
                if cached_result:
                    logger.info(f"Returning cached result for query: {query_text[:50]}...")
                    return cached_result

            # Create query engine with specific parameters
            # Using configurable response mode for performance optimization
            query_engine = self.index.as_query_engine(
                similarity_top_k=top_k or settings.rag_similarity_top_k,
                response_mode=settings.rag_response_mode,  # Configurable: compact, tree_summarize, simple
                verbose=settings.debug_mode,
                streaming=False,
            )

            # If analysis_type is specified, enhance the query with specialized context
            if analysis_type and analysis_type != "standard":
                query_text = self._enhance_query_with_analysis_type(query_text, analysis_type)

            # Aggiungi prompt per rispondere in italiano
            query_text_it = f"Per favore rispondi in italiano. {query_text}"

            # Execute query
            response = query_engine.query(query_text_it)

            # Extract source information
            sources = []
            if hasattr(response, "source_nodes"):
                for node in response.source_nodes:
                    sources.append(
                        {"text": node.node.text[:200] + "...", "score": node.score, "metadata": node.node.metadata}
                    )

            # If specialized analysis is requested, post-process the response
            if analysis_type and analysis_type != "standard":
                response_text = self._apply_specialized_analysis(str(response), sources, query_text, analysis_type)
            else:
                response_text = str(response)

            result = {
                "answer": response_text,
                "unformattedAnswer": clean_markdown(response_text),
                "sources": sources,
                "confidence": sources[0]["score"] if sources else 0,
                "analysis_type": analysis_type or "standard",
            }

            # Cache the result if caching is enabled
            if self.query_cache:
                self.query_cache.set(query_text, top_k, result, analysis_type)

            return result

        except Exception as e:
            logger.error(f"Error querying index: {str(e)}")
            return {"answer": f"Error processing query: {str(e)}", "sources": [], "confidence": 0}

    def query_enhanced(
        self,
        query_text: str,
        top_k: int = 5,
        filters: Optional[dict[str, Any]] = None,
        analysis_type: Optional[str] = None,
        use_reranking: bool = True,
        use_contextual_chunks: bool = True,
        rerank_top_k: int = 10,
    ) -> dict[str, Any]:
        """Enhanced query with reranking and contextual chunks retrieval.

        Args:
            query_text: Query text
            top_k: Final number of results to return
            filters: Optional filters for retrieval
            analysis_type: Type of specialized analysis
            use_reranking: Whether to use CrossEncoder reranking
            use_contextual_chunks: Whether to include contextual chunks
            rerank_top_k: Number of initial results for reranking

        Returns:
            Enhanced query results with improved relevance
        """
        try:
            if not self.index:
                return {"answer": "Nessun documento è stato ancora indicizzato.", "sources": [], "confidence": 0}

            # Check cache first if enabled
            cache_key = f"{query_text}_{top_k}_{analysis_type}_{use_reranking}_{use_contextual_chunks}"
            if self.query_cache:
                cached_result = self.query_cache.get(cache_key, top_k, analysis_type)
                if cached_result:
                    logger.info(f"Returning cached enhanced result for query: {query_text[:50]}...")
                    return cached_result

            # Get more initial results for reranking/contextual enhancement
            initial_top_k = max(rerank_top_k, top_k * 2) if (use_reranking or use_contextual_chunks) else top_k

            # Create query engine
            query_engine = self.index.as_query_engine(
                similarity_top_k=initial_top_k,
                response_mode=settings.rag_response_mode,
                verbose=settings.debug_mode,
                streaming=False,
            )

            # Enhance query with analysis type if specified
            if analysis_type and analysis_type != "standard":
                query_text = self._enhance_query_with_analysis_type(query_text, analysis_type)

            # Add Italian language prompt
            query_text_it = f"Per favore rispondi in italiano. {query_text}"

            # Execute initial query
            response = query_engine.query(query_text_it)

            # Extract initial source information
            initial_sources = []
            if hasattr(response, "source_nodes"):
                for node in response.source_nodes:
                    source_dict = {
                        "id": getattr(node, "node_id", str(hash(node.node.text))),
                        "content": node.node.text,
                        "text": node.node.text[:200] + "...",
                        "score": float(node.score),
                        "metadata": node.node.metadata,
                    }
                    initial_sources.append(source_dict)

            enhanced_sources = initial_sources
            processing_stats = {"initial_sources": len(initial_sources)}

            # Apply reranking if enabled and available
            if use_reranking and self.reranking_service and self.reranking_service.is_available():
                try:
                    reranked_sources = self.reranking_service.rerank_rag_results(
                        query_text, initial_sources, top_k=top_k
                    )
                    enhanced_sources = reranked_sources
                    processing_stats["reranked_sources"] = len(reranked_sources)
                    logger.info(f"Reranked {len(initial_sources)} to {len(reranked_sources)} sources")
                except Exception as e:
                    logger.warning(f"Reranking failed, using original sources: {e}")

            # Apply contextual chunks retrieval if enabled and available
            if use_contextual_chunks and self.contextual_service:
                try:
                    contextual_chunks = self.contextual_service.enhance_retrieval_results(
                        enhanced_sources, self.vector_store, include_metadata=True
                    )

                    # Convert ChunkContext objects back to source format
                    enhanced_sources = []
                    for chunk in contextual_chunks[:top_k]:
                        enhanced_sources.append(
                            {
                                "text": chunk.content[:200] + "...",
                                "score": chunk.score,
                                "metadata": chunk.metadata,
                                "context_type": chunk.context_type,
                                "source_file": chunk.source_file,
                            }
                        )

                    processing_stats["contextual_sources"] = len(contextual_chunks)
                    context_stats = self.contextual_service.get_context_statistics(contextual_chunks)
                    processing_stats["context_stats"] = context_stats

                    logger.info(f"Enhanced with {len(contextual_chunks)} contextual chunks")
                except Exception as e:
                    logger.warning(f"Contextual enhancement failed, using existing sources: {e}")

            # Apply specialized analysis if requested
            if analysis_type and analysis_type != "standard":
                response_text = self._apply_specialized_analysis(
                    str(response), enhanced_sources, query_text, analysis_type
                )
            else:
                response_text = str(response)

            # Calculate enhanced confidence score
            confidence = self._calculate_enhanced_confidence(enhanced_sources)

            result = {
                "answer": response_text,
                "sources": enhanced_sources,
                "confidence": confidence,
                "analysis_type": analysis_type or "standard",
                "enhancement_used": {
                    "reranking": use_reranking and self.reranking_service is not None,
                    "contextual_chunks": use_contextual_chunks and self.contextual_service is not None,
                },
                "processing_stats": processing_stats,
            }

            # Cache the enhanced result
            if self.query_cache:
                self.query_cache.set(cache_key, top_k, result, analysis_type)

            return result

        except Exception as e:
            logger.error(f"Error in enhanced query: {str(e)}")
            # Fallback to standard query
            return self.query(query_text, top_k, filters, analysis_type)

    def _calculate_enhanced_confidence(self, sources: List[dict]) -> float:
        """Calculate confidence score for enhanced results."""
        if not sources:
            return 0.0

        try:
            # Weight scores based on context type if available
            weighted_scores = []
            for source in sources:
                score = source.get("score", 0)
                context_type = source.get("context_type", "original")

                # Original chunks get full weight, context chunks get reduced weight
                weight = 1.0 if context_type == "original" else 0.7
                weighted_scores.append(score * weight)

            # Return weighted average
            return sum(weighted_scores) / len(weighted_scores)

        except Exception as e:
            logger.warning(f"Error calculating enhanced confidence: {e}")
            # Fallback to simple average
            return sum(source.get("score", 0) for source in sources) / len(sources)

    async def enterprise_query(
        self, query_text: str, enable_enterprise_features: bool = True, **kwargs
    ) -> dict[str, Any]:
        """
        Enterprise query with full validation, normalization, and fact table integration.

        Args:
            query_text: The query to process
            enable_enterprise_features: Whether to use enterprise processing pipeline
            **kwargs: Additional parameters for EnterpriseQuery
        """
        if not enable_enterprise_features or not self.enterprise_orchestrator or not ENTERPRISE_AVAILABLE:
            # Fallback to standard query
            return self.query(query_text, **kwargs)

        try:
            # Create enterprise query - filter known parameters
            valid_params = {
                "query_text": query_text,
                "top_k": kwargs.get("top_k"),
                "use_context": kwargs.get("use_context"),
                "csv_analysis": kwargs.get("csv_analysis"),
                "confidence_threshold": kwargs.get("confidence_threshold", 70.0),
            }

            # Remove None values
            valid_params = {k: v for k, v in valid_params.items() if v is not None}

            enterprise_query = EnterpriseQuery(**valid_params)

            # Get current documents for processing
            documents = []
            if hasattr(self, "_last_document_texts"):
                for filename, content in self._last_document_texts.items():
                    documents.append({"filename": filename, "content": content, "hash": str(hash(content))})

            # Process through enterprise pipeline
            processing_result = await self.enterprise_orchestrator.process_enterprise_query(enterprise_query, documents)

            # Generate AI response and get sources
            ai_response = self._generate_ai_response_with_sources(processing_result)

            # Convert processing result to standard query response format
            enterprise_response = {
                "answer": ai_response["answer"],
                "unformattedAnswer": clean_markdown(ai_response["answer"]),
                "sources": ai_response["sources"],
                "confidence": processing_result.confidence_score,
                "analysis_type": "enterprise",
                "enterprise_data": {
                    "source_references": [ref.to_dict() for ref in processing_result.source_refs],
                    "validation_results": [val.to_dict() for val in processing_result.validation_results],
                    "normalized_metrics": {
                        k: {
                            "value": v.to_float(),
                            "original": v.original_value,
                            "scale": v.scale_applied.unit_name,
                            "confidence": v.confidence,
                        }
                        for k, v in processing_result.normalized_data.items()
                    },
                    "mapped_metrics": processing_result.mapped_metrics,
                    "fact_table_records": processing_result.fact_table_records,
                    "processing_time_ms": processing_result.processing_time_ms,
                    "warnings": processing_result.warnings,
                    "errors": processing_result.errors,
                },
            }

            logger.info(f"Enterprise query processed with confidence {processing_result.confidence_score:.1%}")
            return enterprise_response

        except Exception as e:
            logger.error(f"Enterprise query failed, falling back to standard: {e}")
            # Fallback to standard query on error
            return self.query(query_text, **kwargs)

    def _generate_ai_response_with_sources(self, processing_result) -> dict[str, Any]:
        """Generate AI response with sources from processing result."""
        try:
            # Use standard query method to get AI response and sources
            if hasattr(processing_result, "query_text") and self.index:
                standard_response = self.query(processing_result.query_text, top_k=3)

                # Format enterprise response with AI answer + enterprise metadata
                enhanced_answer = standard_response["answer"]

                # Add enterprise metadata to the response
                metadata_parts = []

                # Add normalized metrics if found
                if processing_result.normalized_data:
                    metadata_parts.append("\n\n## 📊 Metriche Finanziarie Rilevate")
                    for metric_name, normalized in processing_result.normalized_data.items():
                        mapped = processing_result.mapped_metrics.get(metric_name)
                        canonical_name = mapped.get("canonical_name", metric_name) if mapped else metric_name
                        metadata_parts.append(
                            f"- **{canonical_name}**: {normalized.to_base_units():,.0f}"
                            f" (scala: {normalized.scale_applied.unit_name})"
                        )

                # Add validation warnings if any
                if processing_result.validation_results:
                    failed_validations = [v for v in processing_result.validation_results if not v.passed]
                    if failed_validations:
                        metadata_parts.append("\n\n## ⚠️ Avvisi di Validazione")
                        for validation in failed_validations:
                            metadata_parts.append(f"- {validation.message}")

                # Add processing metadata at the end
                processing_metadata = []
                if processing_result.fact_table_records > 0:
                    processing_metadata.append(f"{processing_result.fact_table_records} record salvati nel fact table")
                processing_metadata.append(
                    f"Elaborato in {processing_result.processing_time_ms:.0f}ms con confidenza {processing_result.confidence_score:.1%}"
                )

                if processing_metadata:
                    metadata_parts.append(f"\n\n*{' • '.join(processing_metadata)}*")

                # Combine AI answer with enterprise metadata
                if metadata_parts:
                    enhanced_answer = enhanced_answer + "\n".join(metadata_parts)

                return {"answer": enhanced_answer, "sources": standard_response["sources"]}

            # Fallback if no query_text or index
            return {
                "answer": f"Elaborato in {processing_result.processing_time_ms:.0f}ms con confidenza {processing_result.confidence_score:.1%}",
                "sources": [],
            }

        except Exception as e:
            logger.error(f"Error generating AI response with sources: {e}")
            return {"answer": f"Errore nella generazione della risposta: {str(e)}", "sources": []}

    def _enhance_query_with_analysis_type(self, query_text: str, analysis_type: str) -> str:
        """Enhance query based on analysis type."""
        enhancements = {
            "bilancio": """
                        Sei un equity/credit analyst esperto specializzato in documenti finanziari italiani. Analizza il la domanda applicando le seguenti competenze:

                        COMPETENZE SPECIALIZZATE:
                        - **Numeri italiani**: 1.234,56 = milleduecentotrentaquattro virgola cinquantasei
                        - **Negativi**: (123) = numero negativo, -123
                        - **Percentuali**: 5,2% = cinque virgola due per cento
                        - **Scale**: "valori in migliaia" significa moltiplicare × 1.000
                        - **Sinonimi**: fatturato = ricavi = vendite; EBITDA = MOL; PFN = posizione finanziaria netta
                        - **Validazioni**: Attivo = Passivo; PFN = Debito lordo - Cassa; Margine lordo = Ricavi - COGS

                        ISTRUZIONI OPERATIVE:
                        1. **Parsing accurato**: Riconosci formato numerico italiano (es. 1.234.567,89)
                        2. **Provenienza precisa**: Cita sempre "p. X" o "tab. Y" per ogni numero
                        3. **Controlli coerenza**: Verifica equazioni contabili basilari
                        4. **Scale applicate**: Se dichiarato "in migliaia", converti automaticamente
                        5. **Sinonimi**: Normalizza "fatturato" → "ricavi", "MOL" → "EBITDA"



                        REGOLE
                        - Considera metriche come ricavi, EBITDA, margini, cash flow, PFN e covenant.
                        - Fornisci numeri specifici e confronti YoY quando disponibili
                         """,
            "fatturato": """
                        Agisci come sales/revenue analyst esperto. Analizza il la domadna data (no fonti esterne).
                        Focalizzati su vendite e fatturato.
                        Considera ASP, volumi, mix prodotto/cliente, pipeline e forecast.
                        Evidenzia trend di crescita e driver principali
                        Cita "p. X" dopo i numeri
                         """,
            "magazzino": """
                        Agisci come operations/inventory analyst esperto. Analizza la domanda dal punto di vista logistico e di gestione scorte.
                        Focalizzati su rotazione, OTIF, obsoleti, lead time, livelli di servizio, rischi operativi e prossimi passi, con citazioni "p. X".
                         """,
            "contratto": """
                        Agisci come legal/ops analyst esperto. Analizza la domanda dal punto di vista legale e contrattuale.
                        Considera obblighi, SLA, penali, responsabilità, clausole rilevanti e red flag. Cita pagine (p. X).

                        REGOLE: non inferire; se mancano dettagli, fallo notare e non prendere iniziative.
                         """,
            "presentazione": """
                        Sei un business analyst esperto. Analizza la domanda estraendo i messaggi chiave e le raccomandazioni strategiche.
                        Identifica obiettivi, milestone e next steps.
                        Cattura l'essenza della presentazione, i messaggi chiave e le raccomandazioni principali. Cita "slide X" o "p. X" per riferimenti specifici.
                             """,
            "report_dettagliato": """
                                Sei un senior equity research analyst specializzato in documenti finanziari italiani. Produci un'analisi professionale approfondita della domanda in stile investment memorandum.

                                COMPETENZE AVANZATE RICHIESTE:
                                - **Parsing numeri italiani**: 1.234.567,89 (format italiano), (123) = negativo, 5,2%
                                - **Gestione scale**: "valori in migliaia/milioni" → conversione automatica
                                - **Sinonimi finanziari**: fatturato=ricavi=vendite; EBITDA=MOL; PFN=debito netto
                                - **Validazioni contabili**: Attivo=Passivo, PFN=Debito-Cassa, Margine=Ricavi-COGS
                                - **Provenienza granulare**: "p.12|tab.1|riga:Ricavi" per ogni numero
                                - **Confronti strutturati**: YoY%, vs Budget%, scostamenti quantificati

                                METODOLOGIA OPERATIVA:
                                1. **Estrazione accurata**: Riconosci tutti i numeri in formato italiano
                                2. **Conversioni**: Applica scale dichiarate ("in migliaia" × 1.000)
                                3. **Normalizzazione**: Uniforma sinonimi (fatturato → ricavi)
                                4. **Validazione**: Controlla coerenze contabili basilari  
                                5. **Bridge analysis**: Spiega variazioni con numeri precisi
                                6. **Citations**: Ogni dato con fonte esatta (p.X, tab.Y)

                                Includi executive summary, analisi per sezione, KPI, rischi e raccomandazioni.
                                Quantifica sempre le variazioni e usa confronti YoY/Budget.
                                   """,
        }

        enhancement = enhancements.get(analysis_type, "")
        if enhancement:
            return f"{enhancement}\n\nDomanda: {query_text}"
        return query_text

    def _apply_specialized_analysis(self, response: str, sources: list[dict], query: str, analysis_type: str) -> str:
        """Apply specialized post-processing based on analysis type."""
        try:
            from openai import OpenAI

            client = OpenAI(api_key=settings.openai_api_key)

            # Prepare source context
            source_context = "\n\n".join(
                [
                    f"Fonte {i + 1} (rilevanza: {s['score']:.2f}):\n{s['text']}"
                    for i, s in enumerate(sources[:3])  # Use top 3 sources
                ]
            )

            # Create specialized prompt based on analysis type
            specialized_prompts = {
                "bilancio": f"""Riformula questa risposta con focus finanziario professionale:

                            Risposta originale: {response}

                            Fonti rilevanti:
                            {source_context}

                            Produci un'analisi finanziaria strutturata che include:
                            - Metriche chiave con valori specifici
                            - Trend e variazioni percentuali
                            - Confronti con periodi precedenti
                            - Implicazioni finanziarie

                            Usa un tono da equity analyst professionale.""",
                "report_dettagliato": f"""Trasforma questa risposta in un briefing professionale dettagliato:

                            Risposta originale: {response}

                            Fonti rilevanti:
                            {source_context}

                            Struttura l'analisi come:
                            ## Executive Summary
                            [Sintesi in 2-3 righe]

                            ## Analisi Dettagliata
                            [Punti chiave con quantificazione]

                            ## Implicazioni e Raccomandazioni
                            [Azioni suggerite basate sui dati]

                            Mantieni un tono professionale da investment analyst.""",
                "fatturato": f"""Riformula con focus su vendite e revenue:

                            Risposta: {response}
                            Fonti: {source_context}

                            Evidenzia: driver di crescita, mix prodotto, trend vendite, forecast.""",
            }

            prompt = specialized_prompts.get(analysis_type)
            if not prompt:
                return response  # Return original if no specialized prompt

            # Get specialized analysis
            completion = client.chat.completions.create(
                model=settings.llm_model,
                messages=[
                    {
                        "role": "system",
                        "content": "Sei un analista esperto che fornisce analisi specializzate basate sul tipo di documento.",
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.0,
                max_tokens=1000,
            )

            specialized_response = completion.choices[0].message.content

            # Add analysis type marker
            return f"[Analisi {analysis_type.upper()}]\n\n{specialized_response}"

        except Exception as e:
            logger.error(f"Error applying specialized analysis: {str(e)}")
            return response  # Return original response on error

    def query_with_context(
        self, query_text: str, context_data: dict[str, Any], top_k: int = 2, analysis_type: Optional[str] = None
    ) -> dict[str, Any]:
        """Query with additional context from CSV analysis."""
        try:
            # Enhance query with context
            enhanced_query = f"""
            Basandoti sul seguente contesto di dati aziendali:
            {self._format_context(context_data)}

            Domanda: {query_text}

            Per favore fornisci una risposta dettagliata considerando sia i documenti che i dati aziendali forniti. Rispondi in italiano.
            """

            return self.query(enhanced_query, top_k=top_k, analysis_type=analysis_type)

        except Exception as e:
            logger.error(f"Error in query with context: {str(e)}")
            return {"answer": f"Error processing query: {str(e)}", "sources": [], "confidence": 0}

    def _format_context(self, context_data: dict[str, Any]) -> str:
        """Format context data for query enhancement."""
        formatted_parts = []

        if "summary" in context_data:
            formatted_parts.append("Metriche di riepilogo:")
            for key, value in context_data["summary"].items():
                formatted_parts.append(f"- {key}: {value}")

        if "trends" in context_data:
            formatted_parts.append("\nTendenze:")
            if "yoy_growth" in context_data["trends"]:
                for growth in context_data["trends"]["yoy_growth"][-2:]:  # Last 2 years
                    formatted_parts.append(f"- Anno {growth['year']}: {growth['growth_percentage']}% crescita")

        if "insights" in context_data:
            formatted_parts.append("\nApprofondimenti chiave:")
            for insight in context_data["insights"][:3]:  # Top 3 insights
                formatted_parts.append(f"- {insight}")

        return "\n".join(formatted_parts)

    async def query_stream(self, query_text: str, top_k: int = 3, **kwargs):
        """
        Stream query response in real-time.

        Args:
            query_text: User query
            top_k: Number of documents to retrieve
            **kwargs: Additional arguments

        Yields:
            StreamingChunk objects with tokens and metadata
        """
        if not self.streaming_engine:
            logger.warning("Streaming not available, falling back to standard query")
            result = self.query(query_text, top_k=top_k)
            # Convert standard response to single chunk
            if ADVANCED_FEATURES_AVAILABLE:
                yield StreamingChunk(
                    token=result.get("answer", ""), metadata={"sources": result.get("sources", [])}, is_final=True
                )
            else:
                yield {"answer": result.get("answer", ""), "sources": result.get("sources", [])}
            return

        # Stream the response
        async for chunk in self.streaming_engine.stream_query(query_text, **kwargs):
            yield chunk

    def query_with_hyde(self, query_text: str, top_k: int = 3, **kwargs) -> dict[str, Any]:
        """
        Query using HyDE for improved retrieval.

        Args:
            query_text: User query
            top_k: Number of documents to retrieve
            **kwargs: Additional arguments

        Returns:
            Query response with improved retrieval
        """
        if not self.hyde_engine:
            logger.warning("HyDE not available, falling back to standard query")
            return self.query(query_text, top_k=top_k)

        try:
            # Use HyDE engine for query
            response = self.hyde_engine.query(query_text, **kwargs)

            # Format response
            sources = []
            if hasattr(response, "source_nodes"):
                for node in response.source_nodes[:top_k]:
                    sources.append(
                        {
                            "content": node.node.text[:500],
                            "metadata": node.node.metadata,
                            "score": float(node.score) if node.score else 0.0,
                        }
                    )

            result = {
                "answer": str(response),
                "sources": sources,
                "confidence": sum(s["score"] for s in sources) / len(sources) if sources else 0.0,
                "method": "hyde",
            }

            logger.info(f"HyDE query completed with {len(sources)} sources")
            return result

        except Exception as e:
            logger.error(f"Error in HyDE query: {str(e)}")
            return self.query(query_text, top_k=top_k)

    def benchmark_hyde(self, test_queries: list[str]) -> dict[str, float]:
        """
        Benchmark HyDE improvement over standard retrieval.

        Args:
            test_queries: List of test queries

        Returns:
            Benchmark results
        """
        if not self.hyde_engine or not self.index:
            return {"error": "HyDE or index not available"}

        try:
            base_retriever = self.index.as_retriever()
            results = self.hyde_engine.benchmark_improvement(test_queries=test_queries, base_retriever=base_retriever)

            logger.info(f"HyDE Benchmark Results: {results}")
            return results

        except Exception as e:
            logger.error(f"Error benchmarking HyDE: {str(e)}")
            return {"error": str(e)}

    def delete_documents(self, source_filter: str) -> bool:
        """Delete documents by source filter."""
        try:
            # Delete the entire collection and recreate it
            logger.info(f"Deleting collection: {self.collection_name}")

            # Clear query cache when deleting documents
            if self.query_cache:
                self.query_cache.clear()
                logger.info("Query cache cleared after document deletion")

            # Delete the collection completely
            try:
                self.client.delete_collection(self.collection_name)
                logger.info(f"Collection {self.collection_name} deleted successfully")
            except Exception as e:
                logger.warning(f"Collection might not exist: {str(e)}")

            # Recreate the collection from scratch
            vector_size = 1536  # OpenAI text-embedding-3-small dimension
            self.client.create_collection(
                collection_name=self.collection_name,
                vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
            )
            logger.info(f"Created fresh collection: {self.collection_name}")

            # Reinitialize the index
            self._initialize_index()
            logger.info("Successfully cleared all documents from collection")
            return True

        except Exception as e:
            logger.error(f"Error deleting documents: {str(e)}")
            return False

    def get_index_stats(self) -> dict[str, Any]:
        """Get statistics about the indexed documents."""
        try:
            collection_info = self.client.get_collection(self.collection_name)

            return {
                "total_vectors": collection_info.points_count,
                "collection_name": self.collection_name,
                "vector_dimension": collection_info.config.params.vectors.size,
                "distance_metric": str(collection_info.config.params.vectors.distance),
            }

        except Exception as e:
            logger.error(f"Error getting index stats: {str(e)}")
            return {"total_vectors": 0, "error": str(e)}

    def explore_database(
        self,
        limit: int = 100,
        offset: int = 0,
        search_text: Optional[str] = None,
        date_from: Optional[datetime] = None,
        date_to: Optional[datetime] = None,
    ) -> dict[str, Any]:
        """Explore documents in the vector database with pagination, search, and date filtering.

        Args:
            limit: Maximum number of points to retrieve
            offset: Starting offset for pagination
            search_text: Optional text to search for
            date_from: Optional start date filter
            date_to: Optional end date filter

        Returns:
            Dictionary containing documents and statistics
        """
        try:
            # Get collection info
            collection_info = self.client.get_collection(self.collection_name)
            total_points = collection_info.points_count

            if total_points == 0:
                return {"documents": [], "total_count": 0, "unique_sources": [], "stats": {}}

            # Build filter conditions for date range
            filter_conditions = []

            if date_from or date_to:
                # Create date range filter
                date_filter_conditions = []

                if date_from:
                    date_filter_conditions.append(
                        FieldCondition(key="metadata.indexed_at", range=DatetimeRange(gte=date_from.isoformat()))
                    )

                if date_to:
                    date_filter_conditions.append(
                        FieldCondition(key="metadata.indexed_at", range=DatetimeRange(lte=date_to.isoformat()))
                    )

                filter_conditions.extend(date_filter_conditions)

            # Create scroll filter if we have conditions
            scroll_filter = None
            if filter_conditions:
                scroll_filter = Filter(must=filter_conditions)

            # Retrieve points with pagination and filters
            # Note: Qdrant doesn't support traditional pagination, so we use scroll
            # Prende dal db
            points, next_offset = self.client.scroll(
                collection_name=self.collection_name,
                limit=limit,
                offset=offset,
                scroll_filter=scroll_filter,
                with_payload=True,
                with_vectors=False,  # Don't need vectors for exploration
            )

            # Process points to extract document information
            documents = []
            source_set = set()
            file_types = {}
            total_size = 0

            for point in points:
                payload = point.payload if point.payload else {}

                # Extract metadata
                source = payload.get("source", "Unknown")
                source_set.add(source)

                doc_info = {
                    "id": str(point.id),
                    "source": source,
                    "indexed_at": payload.get("indexed_at", "Unknown"),
                    "file_type": payload.get("file_type", "Unknown"),
                    "document_size": payload.get("document_size", 0),
                    "page": payload.get("page", None),
                    "total_pages": payload.get("total_pages", None),
                    "pdf_path": payload.get("pdf_path", None),
                    "text_preview": payload.get("_node_content", "")[:200] + "..."
                    if payload.get("_node_content")
                    else "No content",
                    "metadata": payload,
                }

                documents.append(doc_info)

                # Collect stats
                file_type = doc_info["file_type"]
                file_types[file_type] = file_types.get(file_type, 0) + 1
                total_size += doc_info["document_size"]

            # Get unique sources with more details
            unique_sources = self._get_unique_sources_details()

            return {
                "documents": documents,
                "total_count": total_points,
                "current_page": offset // limit + 1,
                "total_pages": (total_points + limit - 1) // limit,
                "unique_sources": unique_sources,
                "stats": {
                    "total_documents": len(unique_sources),
                    "total_chunks": total_points,
                    "file_types": file_types,
                    "total_size_bytes": total_size,
                    "avg_chunk_size": total_size // total_points if total_points > 0 else 0,
                },
            }

        except Exception as e:
            logger.error(f"Error exploring database: {str(e)}")
            return {"documents": [], "total_count": 0, "unique_sources": [], "stats": {}, "error": str(e)}

    def _get_unique_sources_details(self) -> list[dict[str, Any]]:
        """Get detailed information about unique document sources."""
        try:
            # Use scroll to get all points and extract unique sources
            all_points = []
            offset = None

            while True:
                points, next_offset = self.client.scroll(
                    collection_name=self.collection_name,
                    limit=100,
                    offset=offset,
                    with_payload=True,
                    with_vectors=False,
                )

                if not points:
                    break

                all_points.extend(points)
                offset = next_offset

                if offset is None:
                    break

            # Group by source
            sources_map = {}
            for point in all_points:
                payload = point.payload if point.payload else {}
                source = payload.get("source", "Unknown")

                if source not in sources_map:
                    sources_map[source] = {
                        "name": source,
                        "file_type": payload.get("file_type", "Unknown"),
                        "indexed_at": payload.get("indexed_at", "Unknown"),
                        "chunk_count": 0,
                        "total_size": 0,
                        "pages": set(),
                        "pdf_path": payload.get("pdf_path"),
                        "has_analysis": False,
                    }

                sources_map[source]["chunk_count"] += 1
                sources_map[source]["total_size"] += payload.get("document_size", 0)

                if payload.get("page"):
                    sources_map[source]["pages"].add(payload.get("page"))

            # Convert to list and process
            unique_sources = []
            for source_info in sources_map.values():
                source_info["page_count"] = len(source_info["pages"]) if source_info["pages"] else None
                source_info["pages"] = sorted(source_info["pages"]) if source_info["pages"] else []

                # Check if we have cached analysis for this document
                if hasattr(self, "_last_document_texts") and source_info["name"] in self._last_document_texts:
                    source_info["has_analysis"] = True

                unique_sources.append(source_info)

            # Sort by indexed date (most recent first)
            unique_sources.sort(key=lambda x: x["indexed_at"], reverse=True)

            return unique_sources

        except Exception as e:
            logger.error(f"Error getting unique sources: {str(e)}")
            return []

    def get_document_chunks(self, source_name: str) -> list[dict[str, Any]]:
        """Get all chunks for a specific document source."""
        try:
            from qdrant_client.models import FieldCondition, Filter, MatchValue

            # Create filter for specific source
            filter_condition = Filter(must=[FieldCondition(key="source", match=MatchValue(value=source_name))])

            # Retrieve all chunks for this source
            chunks = []
            offset = None

            while True:
                points, next_offset = self.client.scroll(
                    collection_name=self.collection_name,
                    scroll_filter=filter_condition,
                    limit=100,
                    offset=offset,
                    with_payload=True,
                    with_vectors=False,
                )

                if not points:
                    break

                for point in points:
                    payload = point.payload if point.payload else {}
                    chunk_info = {
                        "id": str(point.id),
                        "page": payload.get("page"),
                        "text": payload.get("_node_content", "")[:500]
                        if payload.get("_node_content")
                        else "No content",
                        "size": payload.get("document_size", 0),
                        "metadata": payload,
                    }
                    chunks.append(chunk_info)

                offset = next_offset
                if offset is None:
                    break

            # Sort by page number if available
            chunks.sort(key=lambda x: x["page"] if x["page"] else 0)

            return chunks

        except Exception as e:
            logger.error(f"Error getting document chunks: {str(e)}")
            return []

    def delete_document_by_source(self, source_name: str) -> bool:
        """Delete all chunks for a specific document source."""
        try:
            from qdrant_client.models import FieldCondition, Filter, MatchValue

            # Create filter for specific source
            filter_condition = Filter(must=[FieldCondition(key="source", match=MatchValue(value=source_name))])

            # Get all point IDs for this source
            points_to_delete = []
            offset = None

            while True:
                points, next_offset = self.client.scroll(
                    collection_name=self.collection_name,
                    scroll_filter=filter_condition,
                    limit=100,
                    offset=offset,
                    with_payload=False,
                    with_vectors=False,
                )

                if not points:
                    break

                points_to_delete.extend([point.id for point in points])

                offset = next_offset
                if offset is None:
                    break

            # Delete points
            if points_to_delete:
                self.client.delete(collection_name=self.collection_name, points_selector=points_to_delete)
                logger.info(f"Deleted {len(points_to_delete)} chunks for document '{source_name}'")

                # Clear from cache if present
                if hasattr(self, "_last_document_texts") and source_name in self._last_document_texts:
                    del self._last_document_texts[source_name]

                # Clear query cache as document changed
                if self.query_cache:
                    self.query_cache.clear()
                    logger.info("Query cache cleared after document deletion")

                return True
            else:
                logger.warning(f"No chunks found for document '{source_name}'")
                return False

        except Exception as e:
            logger.error(f"Error deleting document: {str(e)}")
            return False

    def search_in_database(self, search_query: str, limit: int = 20) -> list[dict[str, Any]]:
        """Search for specific content in the database using semantic search."""
        try:
            # Use the existing query infrastructure for semantic search
            query_engine = self.index.as_query_engine(
                similarity_top_k=limit,
                response_mode="no_text",  # Just get the nodes, no synthesis
            )

            response = query_engine.query(search_query)

            results = []
            if hasattr(response, "source_nodes"):
                for node in response.source_nodes:
                    results.append(
                        {
                            "source": node.node.metadata.get("source", "Unknown"),
                            "page": node.node.metadata.get("page"),
                            "score": node.score,
                            "text": node.node.text[:300] + "..." if len(node.node.text) > 300 else node.node.text,
                            "metadata": node.node.metadata,
                        }
                    )

            return results

        except Exception as e:
            logger.error(f"Error searching in database: {str(e)}")
            return []

    def generate_faq(self, num_questions: int = 10) -> dict[str, Any]:
        """Generate FAQ based on vector database content.

        Args:
            num_questions: Number of FAQ questions to generate

        Returns:
            Dictionary containing generated FAQ with questions and answers
        """
        try:
            if not self.index:
                return {"error": "Nessun documento indicizzato. Carica documenti prima di generare FAQ.", "faqs": []}

            # Get database statistics to understand content
            stats = self.get_index_stats()
            if stats.get("total_vectors", 0) == 0:
                return {"error": "Database vettoriale vuoto. Carica documenti prima di generare FAQ.", "faqs": []}

            logger.info(f"Generating FAQ with {num_questions} questions from {stats.get('total_vectors')} vectors")

            # Get a LARGER sample of documents to understand content themes
            # Increased from 20 to 100 for better content representation
            exploration_data = self.explore_database(limit=100)
            sample_texts = []
            document_types = set()

            if exploration_data.get("documents"):
                for doc in exploration_data["documents"]:
                    # Get actual content, NOT metadata!
                    text_preview = ""

                    # Try to get actual node content first
                    if "metadata" in doc and "_node_content" in doc["metadata"]:
                        text_preview = doc["metadata"]["_node_content"][:1000]
                    elif "text_preview" in doc:
                        text_preview = doc["text_preview"]

                    # Filter out metadata strings and technical info
                    if text_preview and not any(
                        skip in text_preview.lower()
                        for skip in [
                            "metadata",
                            "indexed_at",
                            "source:",
                            "page:",
                            "document_size",
                            "file_type",
                            "pdf_path",
                            "chunk_count",
                            "total_pages",
                        ]
                    ):
                        sample_texts.append(text_preview)
                    if "metadata" in doc:
                        # FIRST: Check if we have explicit document_type metadata (from JSON detection)
                        if "document_type" in doc["metadata"]:
                            doc_type = doc["metadata"]["document_type"]
                            document_types.add(doc_type)
                        # FALLBACK: Check filename-based detection
                        elif "source" in doc["metadata"]:
                            source = doc["metadata"]["source"].lower()
                            if "bilancio" in source or "balance" in source:
                                document_types.add("bilancio")
                            elif "fatturato" in source or "revenue" in source or "vendite" in source:
                                document_types.add("fatturato")
                            elif "contratto" in source or "contract" in source:
                                document_types.add("contratto")
                            elif "report" in source:
                                document_types.add("report")
                            else:
                                document_types.add("generale")

            # Create context for FAQ generation
            # Use MORE samples (30 instead of 10) for better context
            content_sample = "\n\n---\n\n".join(sample_texts[:30]) if sample_texts else ""
            document_types_str = ", ".join(document_types) if document_types else "documenti aziendali"

            # PRIORITIZE: If we have cached full text, use it INSTEAD of chunks
            if hasattr(self, "_last_document_texts") and self._last_document_texts:
                # Use actual full document text for better FAQ generation
                full_texts = list(self._last_document_texts.values())
                if full_texts and any(text.strip() for text in full_texts):
                    # Take substantial portion of actual document content
                    content_sample = "\n\n".join([text[:20000] for text in full_texts])[:30000]
                    logger.info(f"Using FULL document text for FAQ generation: {len(content_sample)} chars")
            elif not sample_texts:
                # If we have no good samples, try a different approach
                logger.warning("No valid text samples found, attempting direct query")
                # Try to retrieve actual content through a query
                try:
                    test_response = self.query("riassumi il contenuto principale del documento", top_k=10)
                    if test_response and "sources" in test_response:
                        for source in test_response["sources"]:
                            if "text" in source:
                                sample_texts.append(source["text"])
                        content_sample = "\n\n".join(sample_texts[:20])
                except:
                    pass

            # Validate we have real content, not metadata
            if not content_sample or len(content_sample) < 100:
                logger.error("Insufficient real content for FAQ generation")
                return {
                    "error": "Contenuto insufficiente per generare FAQ. Assicurati che i documenti siano stati indicizzati correttamente.",
                    "faqs": [],
                    "success": False,
                }

            # Check if content is mostly metadata (bad sign)
            metadata_keywords = ["source:", "indexed_at:", "file_type:", "page:", "metadata", "document_size"]
            metadata_count = sum(1 for keyword in metadata_keywords if keyword in content_sample.lower())
            if metadata_count > 5:
                logger.warning(f"Content seems to contain too much metadata ({metadata_count} occurrences)")

            # Generate FAQ using LLM
            faq_prompt = f"""
Sei un esperto di business intelligence e analisi aziendale. Analizza ATTENTAMENTE il contenuto fornito e genera {num_questions} domande frequenti (FAQ) che siano STRETTAMENTE PERTINENTI al contenuto effettivo del documento.

IMPORTANTE: Le domande DEVONO essere basate ESCLUSIVAMENTE sul contenuto fornito sotto. Non inventare domande su argomenti non presenti nel testo.

TIPI DI DOCUMENTI PRESENTI: {document_types_str}

CONTENUTO EFFETTIVO DEL DOCUMENTO (NON metadata tecnici):
{content_sample}

NOTA CRITICA: Ignora completamente riferimenti a metadata, nomi file, date di indicizzazione, ID documento, numero pagine. Concentrati SOLO sul contenuto business del documento!

ISTRUZIONI RIGOROSE:
1. Genera {num_questions} domande basate ESCLUSIVAMENTE sul contenuto fornito sopra
2. NON inventare domande su argomenti NON menzionati nel testo (es. se non ci sono dati su marketing, NON fare domande sul marketing)
3. Le domande devono riferirsi a informazioni EFFETTIVAMENTE PRESENTI nel documento
4. Se il documento è un bilancio, fai domande su voci di bilancio, patrimonio, risultati economici, etc.
5. Usa terminologia italiana e professionale appropriata al tipo di documento
6. Prima di generare ogni domanda, verifica che l'argomento sia PRESENTE nel contenuto fornito

FORMATO RICHIESTO:
Per ogni domanda, fornisci SOLO:
- Una domanda chiara e specifica
- NON fornire risposte, solo le domande

Esempi di domande SOLO se questi argomenti sono presenti nel testo:
- Se c'è un bilancio: "Qual è il valore del patrimonio netto?", "Come sono variati i ricavi?"
- Se ci sono dati vendite: "Quali sono i trend di vendita?"
- Se ci sono dati su clienti: "Qual è il tasso di attrito dei clienti?"

RICORDA: Genera domande SOLO su ciò che è EFFETTIVAMENTE presente nel contenuto fornito!

Genera SOLO le domande, una per riga, numerate da 1 a {num_questions}:
"""

            # Query the LLM for FAQ generation
            llm_response = Settings.llm.complete(faq_prompt)
            faq_questions_text = str(llm_response)

            # Parse questions from response
            questions = []
            for line in faq_questions_text.split("\n"):
                line = line.strip()
                if line and (line[0].isdigit() or line.startswith("-") or line.startswith("•")):
                    # Clean up the question
                    question = line
                    # Remove numbering (1., 2., -, •, etc.)
                    question = question.lstrip("0123456789.-• \t")
                    if question:
                        questions.append(question)

            # Generate answers for each question using RAG
            faqs = []
            for i, question in enumerate(questions[:num_questions], 1):
                try:
                    # Query the database for answer
                    rag_response = self.query(question, top_k=2)

                    answer = rag_response.get(
                        "answer",
                        "Non sono riuscito a trovare informazioni specifiche su questo argomento nei documenti caricati.",
                    )
                    sources = rag_response.get("sources", [])

                    faqs.append(
                        {
                            "id": i,
                            "question": question,
                            "answer": answer,
                            "sources": sources,
                            "generated_at": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                        }
                    )

                except Exception as e:
                    logger.warning(f"Error generating answer for question {i}: {e}")
                    faqs.append(
                        {
                            "id": i,
                            "question": question,
                            "answer": "Errore nella generazione della risposta. Riprova più tardi.",
                            "sources": [],
                            "generated_at": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                        }
                    )

            result = {
                "faqs": faqs,
                "metadata": {
                    "total_questions": len(faqs),
                    "generated_at": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                    "document_types": list(document_types),
                    "total_documents": stats.get("total_vectors", 0),
                },
                "success": True,
            }

            logger.info(f"Successfully generated {len(faqs)} FAQ questions")
            return result

        except Exception as e:
            logger.error(f"Error generating FAQ: {str(e)}")
            return {"error": f"Errore nella generazione delle FAQ: {str(e)}", "faqs": [], "success": False}

    def _load_csv_with_analysis(
        self, file_path: str, csv_analyzer, metadata: Optional[dict[str, Any]] = None
    ) -> list[Document]:
        """Load CSV file with automatic analysis and insights generation."""
        from pathlib import Path

        import pandas as pd

        try:
            # Read CSV
            df = pd.read_csv(file_path)
            file_name = Path(file_path).name

            # Analyze CSV with csv_analyzer
            analysis = csv_analyzer.analyze(df)
            insights = csv_analyzer.generate_insights(df, analysis)

            documents = []

            # 1. Create metadata document
            metadata_text = f"""
            Dataset: {file_name}
            Righe: {len(df)}
            Colonne: {len(df.columns)}
            Colonne disponibili: {", ".join(df.columns.tolist())}

            Tipi di Dati:
            {chr(10).join([f"- {col}: {dtype}" for col, dtype in analysis.get("data_types", {}).items()])}

            Valori Mancanti:
            {chr(10).join([f"- {col}: {count}" for col, count in analysis.get("missing_values", {}).items()]) if "missing_values" in analysis else "Nessun valore mancante"}

            Statistiche Principali:
            {analysis.get("summary_stats", pd.DataFrame()).to_string() if "summary_stats" in analysis else "Non disponibili"}
            """

            doc_metadata = {"source": file_name, "type": "csv_metadata", "file_type": ".csv"}
            if metadata:
                doc_metadata.update(metadata)

            documents.append(Document(text=metadata_text.strip(), metadata=doc_metadata))

            # 2. Create documents from insights
            for i, insight in enumerate(insights[:10], 1):  # Limit to top 10 insights
                insight_doc = Document(
                    text=f"Insight #{i}: {insight}",
                    metadata={
                        "source": file_name,
                        "type": "csv_insight",
                        "insight_id": i,
                        "file_type": ".csv",
                        **(metadata or {}),
                    },
                )
                documents.append(insight_doc)

            # 3. Create documents from sample data (first 100 rows max)
            sample_size = min(100, len(df))
            sample_df = df.head(sample_size)

            # Convert to readable text chunks (10 rows at a time)
            for i in range(0, sample_size, 10):
                chunk = sample_df.iloc[i : i + 10]
                chunk_text = (
                    f"Dati dal CSV '{file_name}' (righe {i + 1}-{min(i + 10, sample_size)}):\n\n{chunk.to_string()}"
                )

                chunk_doc = Document(
                    text=chunk_text,
                    metadata={
                        "source": file_name,
                        "type": "csv_data",
                        "rows_range": f"{i + 1}-{min(i + 10, sample_size)}",
                        "file_type": ".csv",
                        **(metadata or {}),
                    },
                )
                documents.append(chunk_doc)

            logger.info(
                f"CSV '{file_name}' processed: {len(insights)} insights, {len(documents) - 1 - len(insights)} data chunks"
            )
            return documents

        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            # Fallback to basic CSV loading
            return self._load_csv_basic(file_path, metadata)

    def _load_csv_basic(self, file_path: str, metadata: Optional[dict[str, Any]] = None) -> list[Document]:
        """Basic CSV loading without analysis (fallback method)."""
        from pathlib import Path

        import pandas as pd

        try:
            # Read CSV
            df = pd.read_csv(file_path)
            file_name = Path(file_path).name

            # Create a single document with CSV overview
            csv_text = f"""
            File CSV: {file_name}
            Righe: {len(df)}
            Colonne: {len(df.columns)}

            Colonne disponibili:
            {", ".join(df.columns.tolist())}

            Prime 20 righe di dati:
            {df.head(20).to_string()}

            Ultime 5 righe di dati:
            {df.tail(5).to_string()}
            """

            doc_metadata = {"source": file_name, "type": "csv_basic", "file_type": ".csv"}
            if metadata:
                doc_metadata.update(metadata)

            document = Document(text=csv_text.strip(), metadata=doc_metadata)

            logger.info(f"CSV '{file_name}' loaded with basic processing")
            return [document]

        except Exception as e:
            logger.error(f"Error loading CSV {file_path}: {e}")
            # Ultimate fallback - treat as text file
            return self._load_text(file_path, metadata)

    def _load_excel(self, file_path: str, metadata: Optional[dict[str, Any]] = None) -> list[Document]:
        """Load Excel file and convert to documents."""
        from pathlib import Path

        import pandas as pd

        try:
            file_name = Path(file_path).name
            documents = []

            # Try to use the custom Excel parser if available
            try:
                from src.application.parsers.excel_parser import ExcelParser

                parser = ExcelParser()
                extracted_data = parser.parse(file_path)

                # Create document from parsed data
                excel_text = f"""
                File Excel: {file_name}
                Fogli: {len(extracted_data.workbook_metadata.sheets)}
                Autore: {extracted_data.workbook_metadata.author or "N/A"}
                Ultima modifica: {extracted_data.workbook_metadata.modified or "N/A"}
                
                """

                # Add information about each sheet
                for sheet in extracted_data.workbook_metadata.sheets:
                    excel_text += f"\nFoglio: {sheet.name}\n"
                    excel_text += f"Righe: {sheet.max_row}, Colonne: {sheet.max_column}\n"
                    if sheet.tables:
                        excel_text += f"Tabelle trovate: {len(sheet.tables)}\n"
                        for table in sheet.tables:
                            excel_text += f"  - Tabella {table.start_cell}:{table.end_cell} con headers: {', '.join(table.headers[:5])}\n"

                # Add parsed data from dataframes
                for sheet_name, df in extracted_data.data_frames.items():
                    if df is not None and not df.empty:
                        excel_text += f"\n--- Dati dal foglio {sheet_name} ---\n"
                        excel_text += f"Dimensioni: {len(df)} righe x {len(df.columns)} colonne\n"
                        excel_text += df.to_string()  # Load ALL rows, not just head(20)
                        excel_text += "\n"

                doc_metadata = {
                    "source": file_name,
                    "type": "excel_parsed",
                    "file_type": Path(file_path).suffix.lower(),
                    "sheets_count": len(extracted_data.workbook_metadata.sheets),
                }
                if metadata:
                    doc_metadata.update(metadata)

                documents.append(Document(text=excel_text.strip(), metadata=doc_metadata))
                logger.info(f"Excel '{file_name}' loaded with custom parser")

            except (ImportError, Exception) as e:
                # Fallback to pandas if custom parser not available or fails
                logger.info(f"Custom Excel parser failed ({str(e)}), using pandas fallback")

                # Initialize documents list for pandas fallback
                documents = []

                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                excel_text = f"""
                File Excel: {file_name}
                Fogli: {len(excel_file.sheet_names)}
                Fogli disponibili: {", ".join(excel_file.sheet_names)}
                
                """

                # Process each sheet - ALL sheets, not just first 5
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    excel_text += f"\n--- Foglio: {sheet_name} ---\n"
                    excel_text += f"Righe: {len(df)}, Colonne: {len(df.columns)}\n"
                    # Convert column names to strings to handle numeric columns
                    col_names = [str(col) for col in df.columns]
                    excel_text += f"Colonne: {', '.join(col_names)}\n\n"

                    # Add data - if small dataset, include all; if large, create chunks
                    if len(df) <= 100:
                        # For small datasets (<=100 rows), include all data
                        excel_text += "Dati completi:\n"
                        excel_text += df.to_string()
                        excel_text += "\n\n"
                    else:
                        # For larger datasets, split into chunks for better indexing
                        excel_text += f"Dataset grande ({len(df)} righe) - Dividendo in chunks...\n\n"

                        # Add first chunk with ALL data from first chunk, not just head(50)
                        chunk_size = 100
                        excel_text += f"Prime {chunk_size} righe:\n"
                        excel_text += df.head(chunk_size).to_string()
                        excel_text += "\n\n"

                        # Create additional documents for remaining data
                        for i in range(chunk_size, len(df), chunk_size):
                            chunk_end = min(i + chunk_size, len(df))
                            chunk_df = df.iloc[i:chunk_end]

                            chunk_text = f"""
                            File Excel: {file_name}
                            Foglio: {sheet_name}
                            Righe {i + 1} a {chunk_end} di {len(df)}
                            
                            Dati:
                            {chunk_df.to_string()}
                            """

                            chunk_metadata = {
                                "source": file_name,
                                "type": "excel_chunk",
                                "file_type": Path(file_path).suffix.lower(),
                                "sheet_name": sheet_name,
                                "chunk_start": i + 1,
                                "chunk_end": chunk_end,
                                "total_rows": len(df),
                            }
                            if metadata:
                                chunk_metadata.update(metadata)

                            documents.append(Document(text=chunk_text.strip(), metadata=chunk_metadata))

                        excel_text += f"Creati {len(range(chunk_size, len(df), chunk_size))} documenti aggiuntivi per i dati rimanenti.\n"

                doc_metadata = {
                    "source": file_name,
                    "type": "excel_basic",
                    "file_type": Path(file_path).suffix.lower(),
                    "sheets_count": len(excel_file.sheet_names),
                }
                if metadata:
                    doc_metadata.update(metadata)

                documents.append(Document(text=excel_text.strip(), metadata=doc_metadata))
                logger.info(f"Excel '{file_name}' loaded with pandas")

            return documents

        except Exception as e:
            logger.error(f"Error loading Excel {file_path}: {e}")
            # Fallback - treat as text file
            return self._load_text(file_path, metadata)

    def _load_image(self, file_path: str, metadata: Optional[dict[str, Any]] = None) -> list[Document]:
        """Load image file and convert to text using OCR."""
        from pathlib import Path

        try:
            file_name = Path(file_path).name

            # Use the existing image parser
            try:
                from src.application.services.format_parsers import ImageParser

                parser = ImageParser(ocr_language="ita+eng")  # Italian + English

                # Parse image with OCR
                parsed_content = parser.parse(file_path)

                # Extract OCR text
                ocr_text = parsed_content.data.get("text", "").strip()

                if not ocr_text:
                    ocr_text = f"[IMMAGINE] Il file '{file_name}' non contiene testo riconoscibile tramite OCR."

                # Create document with OCR text
                image_text = f"""
                Immagine: {file_name}
                Formato: {parsed_content.metadata.get("format", "N/A")}
                Dimensioni: {parsed_content.metadata.get("width", 0)}x{parsed_content.metadata.get("height", 0)} px
                
                Testo estratto con OCR:
                {ocr_text}
                """

                # Add table information if found
                if parsed_content.tables:
                    image_text += f"\n\nTabelle rilevate: {len(parsed_content.tables)}\n"
                    for i, table in enumerate(parsed_content.tables[:3], 1):  # Max 3 tables
                        image_text += f"Tabella {i}: {table.get('summary', 'Dati tabulari rilevati')}\n"

                doc_metadata = {
                    "source": file_name,
                    "type": "image_ocr",
                    "file_type": Path(file_path).suffix.lower(),
                    "image_format": parsed_content.metadata.get("format"),
                    "image_size": f"{parsed_content.metadata.get('width', 0)}x{parsed_content.metadata.get('height', 0)}",
                    "ocr_language": "ita+eng",
                    "extraction_time": parsed_content.extraction_time,
                }
                if metadata:
                    doc_metadata.update(metadata)

                document = Document(text=image_text.strip(), metadata=doc_metadata)

                logger.info(f"Image '{file_name}' processed with OCR: {len(ocr_text)} characters extracted")
                return [document]

            except ImportError:
                # Fallback to basic OCR if custom parser not available
                logger.info("Custom image parser not available, using basic pytesseract")

                # Configure Tesseract path for Windows
                import os

                from PIL import Image
                import pytesseract

                if os.name == "nt":  # Windows
                    tesseract_paths = [
                        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                        r"C:\Tesseract-OCR\tesseract.exe",
                    ]
                    for path in tesseract_paths:
                        if os.path.exists(path):
                            pytesseract.pytesseract.tesseract_cmd = path
                            break

                # Open image
                image = Image.open(file_path)

                # Extract text using OCR
                ocr_text = pytesseract.image_to_string(image, lang="ita+eng")

                if not ocr_text.strip():
                    ocr_text = f"[IMMAGINE] Il file '{file_name}' non contiene testo riconoscibile tramite OCR."

                image_text = f"""
                Immagine: {file_name}
                Formato: {image.format}
                Dimensioni: {image.width}x{image.height} px
                
                Testo estratto con OCR:
                {ocr_text}
                """

                doc_metadata = {
                    "source": file_name,
                    "type": "image_ocr_basic",
                    "file_type": Path(file_path).suffix.lower(),
                    "image_format": image.format,
                    "image_size": f"{image.width}x{image.height}",
                    "ocr_language": "ita+eng",
                }
                if metadata:
                    doc_metadata.update(metadata)

                document = Document(text=image_text.strip(), metadata=doc_metadata)

                logger.info(f"Image '{file_name}' processed with basic OCR: {len(ocr_text)} characters extracted")
                return [document]

        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            # Fallback - create document with error info
            return [
                Document(
                    text=f"Errore nel processamento dell'immagine '{Path(file_path).name}': {str(e)}",
                    metadata={"source": Path(file_path).name, "type": "image_error", "error": str(e)},
                )
            ]

    def get_embeddings_statistics(self, sample_size: int = 100) -> dict[str, Any]:
        """Get real embedding statistics from Qdrant.

        Args:
            sample_size: Number of vectors to sample for statistics

        Returns:
            Dictionary with embedding statistics
        """
        try:
            # Get sample points with vectors
            result = self.client.scroll(
                collection_name=self.collection_name, limit=min(sample_size, 1000), with_vectors=True
            )

            if not result or not result[0]:
                return {
                    "error": "No vectors found",
                    "mean": 0,
                    "std": 0,
                    "sparsity": 0,
                    "density": 0,
                    "dimension": settings.embedding_dimension,
                }

            points = result[0]
            vectors = np.array([point.vector for point in points if point.vector])

            if len(vectors) == 0:
                return {
                    "error": "No valid vectors found",
                    "mean": 0,
                    "std": 0,
                    "sparsity": 0,
                    "density": 0,
                    "dimension": settings.embedding_dimension,
                }

            # Calculate statistics
            mean_val = float(np.mean(vectors))
            std_val = float(np.std(vectors))
            sparsity = float(np.mean(vectors == 0))
            density = float(np.mean(vectors != 0))

            # Additional statistics
            min_val = float(np.min(vectors))
            max_val = float(np.max(vectors))
            median_val = float(np.median(vectors))

            return {
                "mean": mean_val,
                "std": std_val,
                "sparsity": sparsity,
                "density": density,
                "min": min_val,
                "max": max_val,
                "median": median_val,
                "dimension": vectors.shape[1] if len(vectors.shape) > 1 else settings.embedding_dimension,
                "sample_size": len(vectors),
                "distribution_data": vectors.flatten().tolist()[:1000],  # First 1000 values for histogram
            }

        except Exception as e:
            logger.error(f"Error getting embeddings statistics: {str(e)}")
            return {
                "error": str(e),
                "mean": 0,
                "std": 0,
                "sparsity": 0,
                "density": 0,
                "dimension": settings.embedding_dimension,
            }

    def calculate_document_similarity(self, doc1_name: str, doc2_name: str) -> float:
        """Calculate cosine similarity between two documents using their embeddings.

        Args:
            doc1_name: Name of first document
            doc2_name: Name of second document

        Returns:
            Cosine similarity score between 0 and 1
        """
        try:
            # Get chunks for both documents
            doc1_vectors = self._get_document_embeddings(doc1_name)
            doc2_vectors = self._get_document_embeddings(doc2_name)

            if doc1_vectors is None or doc2_vectors is None:
                return 0.0

            # Average the embeddings for each document
            doc1_avg = np.mean(doc1_vectors, axis=0)
            doc2_avg = np.mean(doc2_vectors, axis=0)

            # Calculate cosine similarity
            dot_product = np.dot(doc1_avg, doc2_avg)
            norm1 = np.linalg.norm(doc1_avg)
            norm2 = np.linalg.norm(doc2_avg)

            if norm1 == 0 or norm2 == 0:
                return 0.0

            similarity = dot_product / (norm1 * norm2)
            return float(similarity)

        except Exception as e:
            logger.error(f"Error calculating document similarity: {str(e)}")
            return 0.0

    def _get_document_embeddings(self, doc_name: str) -> Optional[np.ndarray]:
        """Get all embeddings for a document.

        Args:
            doc_name: Document name

        Returns:
            Numpy array of embeddings or None
        """
        try:
            # Search for all points belonging to this document
            result = self.client.scroll(
                collection_name=self.collection_name,
                scroll_filter=Filter(must=[FieldCondition(key="metadata.source", match={"value": doc_name})]),
                limit=1000,
                with_vectors=True,
            )

            if not result or not result[0]:
                return None

            points = result[0]
            vectors = [point.vector for point in points if point.vector]

            if not vectors:
                return None

            return np.array(vectors)

        except Exception as e:
            logger.error(f"Error getting document embeddings: {str(e)}")
            return None

    def find_similar_documents(self, source_name: str, top_k: int = 5) -> List[Tuple[str, float]]:
        """Find documents most similar to the given document.

        Args:
            source_name: Name of the source document
            top_k: Number of similar documents to return

        Returns:
            List of tuples (document_name, similarity_score)
        """
        try:
            # Get embeddings for source document
            source_vectors = self._get_document_embeddings(source_name)

            if source_vectors is None:
                return []

            # Average the embeddings
            avg_embedding = np.mean(source_vectors, axis=0)

            # Search for similar vectors
            search_result = self.client.search(
                collection_name=self.collection_name,
                query_vector=avg_embedding.tolist(),
                limit=top_k * 10,  # Get more to filter out same document
            )

            # Group by document and calculate similarities
            doc_scores = {}
            for point in search_result:
                if point.payload and "metadata" in point.payload:
                    doc_name = point.payload["metadata"].get("source", "Unknown")
                    if doc_name != source_name:  # Exclude the source document itself
                        if doc_name not in doc_scores:
                            doc_scores[doc_name] = []
                        doc_scores[doc_name].append(point.score)

            # Average scores per document
            doc_avg_scores = [(doc, np.mean(scores)) for doc, scores in doc_scores.items()]

            # Sort by score and return top_k
            doc_avg_scores.sort(key=lambda x: x[1], reverse=True)
            return doc_avg_scores[:top_k]

        except Exception as e:
            logger.error(f"Error finding similar documents: {str(e)}")
            return []

    def get_document_similarity_matrix(
        self, doc_names: Optional[List[str]] = None, max_docs: int = 10
    ) -> Tuple[np.ndarray, List[str]]:
        """Calculate similarity matrix between documents.

        Args:
            doc_names: List of document names (if None, use first max_docs)
            max_docs: Maximum number of documents to compare

        Returns:
            Tuple of (similarity_matrix, document_names)
        """
        try:
            # Get document list if not provided
            if doc_names is None:
                exploration = self.explore_database(limit=max_docs)
                if not exploration.get("unique_sources"):
                    return np.array([]), []
                doc_names = [doc["name"] for doc in exploration["unique_sources"][:max_docs]]

            n = len(doc_names)
            similarity_matrix = np.zeros((n, n))

            # Calculate pairwise similarities
            for i in range(n):
                for j in range(i, n):
                    if i == j:
                        similarity_matrix[i, j] = 1.0
                    else:
                        sim = self.calculate_document_similarity(doc_names[i], doc_names[j])
                        similarity_matrix[i, j] = sim
                        similarity_matrix[j, i] = sim  # Matrix is symmetric

            return similarity_matrix, doc_names

        except Exception as e:
            logger.error(f"Error creating similarity matrix: {str(e)}")
            return np.array([]), []

    def reindex_documents_batch(
        self, source_names: List[str], force_reindex: bool = True, progress_callback=None
    ) -> dict[str, Any]:
        """Reindex multiple documents in batch.

        Args:
            source_names: List of document source names to reindex
            force_reindex: Whether to force reindexing even if document exists
            progress_callback: Optional callback function for progress updates

        Returns:
            Dictionary with success and failed documents
        """
        results = {"success": [], "failed": [], "skipped": [], "total": len(source_names)}

        for idx, source_name in enumerate(source_names):
            try:
                if progress_callback:
                    progress_callback(idx / len(source_names), f"Processing {source_name}")

                # Check if document exists
                existing_chunks = self.get_document_chunks(source_name)

                if not existing_chunks and not force_reindex:
                    results["skipped"].append({"name": source_name, "reason": "Document not found"})
                    continue

                # If document exists and force_reindex is True, delete it first
                if existing_chunks and force_reindex:
                    # Backup metadata from first chunk
                    metadata_backup = existing_chunks[0].get("metadata", {}) if existing_chunks else {}

                    # Delete existing document
                    delete_success = self.delete_document_by_source(source_name)

                    if not delete_success:
                        results["failed"].append({"name": source_name, "error": "Failed to delete existing document"})
                        continue

                    logger.info(f"Deleted existing document: {source_name}")

                # Find the original file path (if available)
                # This would need to be stored in metadata or provided separately
                # For now, we'll mark it as needing manual re-upload
                results["success"].append(
                    {
                        "name": source_name,
                        "action": "deleted_for_reindex",
                        "message": "Document deleted. Please re-upload the file.",
                    }
                )

            except Exception as e:
                logger.error(f"Error reindexing {source_name}: {str(e)}")
                results["failed"].append({"name": source_name, "error": str(e)})

        if progress_callback:
            progress_callback(1.0, "Reindexing complete")

        return results

    def clear_index(self) -> dict[str, Any]:
        """remove all documents/nodes from the current tenant's vector db"""
        try:
            if not self.index:
                return {"success": False, "message": "No index intialized"}
            deleted = 0
            errors = []
            try:
                stats = self.get_index_stats()
                total_vectors = stats.get("total_vectors", 0)
            except Exception as e:
                logger.warning(f"Could not get index stats:{e}")
                total_vectors = 0
            if total_vectors == 0:
                return {"success": True, "message": "Index already empty", "deleted": 0}

            exploration_data = self.explore_database(limit=total_vectors)
            docs = exploration_data.get("documents", [])
            for doc in docs:
                doc_id = None
                if "metadata" in doc and "doc_id" in doc["metadata"]:
                    doc_id = doc["metadata"]["doc_id"]
                if doc_id:
                    try:
                        self.index.delete_ref_doc(ref_doc_id=doc_id)
                        deleted += 1
                    except Exception as e:
                        errors.append(f"Failed to delete {doc_id}:{e}")
                else:
                    errors.append(f"No doc_id found in document:{doc}")
            logger.info("Vector index cleared successfully")
            return {
                "success": True,
                "message": f"Deleted {deleted} documents" + (f" with {len(errors)} errors" if errors else ""),
                "deleted": deleted,
                "errors": errors,
            }
        except Exception as e:
            logger.error(f"Failed to clear index: {str(e)}")
            return {"success": False, "message": str(e)}
